"""
docx_generator.py — python-docx 기반 DOCX 생성기
FPOF Document Converter
"""
from __future__ import annotations
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..core.content_model import (
    DocumentContent, Section, Block, BlockType, BrandTheme, TableRow
)


def _rgb(hex_str: str):
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_paragraph_color(para, hex_color: str):
    """단락의 모든 run 색상 설정"""
    for run in para.runs:
        run.font.color.rgb = _rgb(hex_color)


def _add_cover(doc: Document, content: DocumentContent, brand: BrandTheme):
    """커버 페이지"""
    doc.add_paragraph()  # 상단 여백

    # 브랜드명
    p = doc.add_paragraph(brand.brand_name.upper())
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.size = Pt(11)
    run.font.color.rgb = _rgb(brand.primary)
    run.font.bold = False

    # 제목
    p = doc.add_paragraph(content.title)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _rgb(brand.secondary)

    # 서브타이틀
    if content.subtitle:
        p = doc.add_paragraph(content.subtitle)
        run = p.runs[0]
        run.font.size = Pt(14)
        run.font.color.rgb = _rgb(brand.text_light)

    # 메타 (시즌, 날짜)
    meta = "  ·  ".join(filter(None, [content.season, str(content.date)]))
    if meta:
        p = doc.add_paragraph(meta)
        run = p.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = _rgb(brand.text_light)

    doc.add_page_break()


def _apply_heading_style(para, level: int, brand: BrandTheme):
    sizes = {1: 22, 2: 16, 3: 13}
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True
    if level == 1:
        run.font.color.rgb = _rgb(brand.secondary)
    elif level == 2:
        run.font.color.rgb = _rgb(brand.primary)
    else:
        run.font.color.rgb = _rgb(brand.text)


def _add_block(doc: Document, block: Block, brand: BrandTheme):
    if block.type == BlockType.HEADING3:
        p = doc.add_paragraph(block.text)
        _apply_heading_style(p, 3, brand)

    elif block.type == BlockType.PARAGRAPH:
        if block.text.strip():
            p = doc.add_paragraph(block.text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if p.runs:
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.color.rgb = _rgb(brand.text)

    elif block.type == BlockType.BULLET:
        for item in block.items:
            p = doc.add_paragraph(f"• {item}", style='Normal')
            p.paragraph_format.left_indent = Inches(0.3)
            if p.runs:
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.color.rgb = _rgb(brand.text)

    elif block.type == BlockType.TABLE:
        _add_table(doc, block, brand)

    elif block.type == BlockType.KPI:
        p = doc.add_paragraph()
        run_val = p.add_run(f"  {block.kpi_value}  ")
        run_val.font.size = Pt(18)
        run_val.font.bold = True
        run_val.font.color.rgb = _rgb(brand.primary)
        if block.kpi_label:
            run_lbl = p.add_run(f"  {block.kpi_label}")
            run_lbl.font.size = Pt(11)
            run_lbl.font.color.rgb = _rgb(brand.text_light)

    elif block.type == BlockType.CODE:
        p = doc.add_paragraph(block.text)
        if p.runs:
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(9)

    elif block.type == BlockType.HR:
        doc.add_paragraph("─" * 40)

    else:
        if block.text.strip():
            doc.add_paragraph(block.text)


def _add_table(doc: Document, block: Block, brand: BrandTheme):
    rows = [r for r in block.rows if r.cells]
    if not rows:
        return
    cols = max(len(r.cells) for r in rows)

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row.cells[:cols]):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(cell_text)

            if row.is_header:
                # 헤더 배경색
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), brand.secondary.lstrip('#'))
                tcPr.append(shd)
                if cell.paragraphs:
                    run = cell.paragraphs[0].add_run()
                    for r in cell.paragraphs[0].runs:
                        r.font.color.rgb = _rgb(brand.background)
                        r.font.bold = True
                        r.font.size = Pt(10)
            else:
                if cell.paragraphs:
                    for r in cell.paragraphs[0].runs:
                        r.font.size = Pt(10)


def generate(doc: DocumentContent, brand: BrandTheme,
             template: str = "internal",
             output_path: str = None) -> str:
    """
    DocumentContent + BrandTheme → DOCX 파일 생성
    """
    document = Document()

    # 페이지 여백
    from docx.shared import Cm
    for section in document.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # 커버
    _add_cover(document, doc, brand)

    # 본문
    for sec in doc.sections:
        if sec.title and sec.title != "(intro)":
            p = document.add_heading(sec.title, level=2)
            _apply_heading_style(p, 2, brand)

        for block in sec.blocks:
            _add_block(document, block, brand)

    # 저장
    if output_path is None:
        base = os.path.splitext(os.path.basename(doc.source_path))[0]
        output_path = os.path.join(
            os.path.dirname(doc.source_path), "exports",
            f"{base}_{template}.docx"
        )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)
    return output_path
